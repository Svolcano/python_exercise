from docx import Document
from PIL import Image
from PIL import ImageDraw
from io import BytesIO


def write_img():
    document = Document()
    p = document.add_paragraph()
    r = p.add_run()
    image_size = 20
    for x in range(10):
        im = Image.new("RGB", (image_size, image_size), 'white')
        draw_obj = ImageDraw.Draw(im)
        draw_obj.ellipse((0, 0, image_size-1, image_size-1), fill=255-x)
        fake_buf_file = BytesIO()
        im.save(fake_buf_file, "png")
        r.add_picture(fake_buf_file)
        fake_buf_file.close()
    document.save("demo.docx")


def read_word():
    document = Document("test.docx")
    for p in document.paragraphs:
        print(p.text)


if __name__ == "__main__":
    read_word()